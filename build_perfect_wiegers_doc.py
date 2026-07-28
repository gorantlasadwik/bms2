import re
import os
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

SRC_FILE = r"C:\Users\sadwi\Downloads\Entity Relationship (ER) Diagram_BACKUP.docx"
DST_FILE = r"C:\Users\sadwi\Downloads\Software_Requirements_Specification_CourtSetu_AI.docx"
DST_FILE2 = r"C:\Users\sadwi\Downloads\Entity_Relationship_Diagram_Wiegers_Formatted.docx"

print("Loading original document...")
src_doc = docx.Document(SRC_FILE)

# Helper function to style tables like Karl Wiegers template
def style_wiegers_table(table):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    tblPr = table._tbl.tblPr
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'<w:top w:val="single" w:sz="12" w:space="0" w:color="000000"/>'
        f'<w:bottom w:val="single" w:sz="12" w:space="0" w:color="000000"/>'
        f'<w:insideH w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        f'<w:insideV w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        f'<w:left w:val="single" w:sz="12" w:space="0" w:color="000000"/>'
        f'<w:right w:val="single" w:sz="12" w:space="0" w:color="000000"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)

    for i, row in enumerate(table.rows):
        is_header = (i == 0)
        for cell in row.cells:
            tcPr = cell._tc.get_or_add_tcPr()
            tcMar = parse_xml(
                f'<w:tcMar {nsdecls("w")}>'
                f'<w:top w:w="100" w:type="dxa"/>'
                f'<w:bottom w:w="100" w:type="dxa"/>'
                f'<w:left w:w="150" w:type="dxa"/>'
                f'<w:right w:w="150" w:type="dxa"/>'
                f'</w:tcMar>'
            )
            tcPr.append(tcMar)

            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(2)
                p.paragraph_format.space_before = Pt(2)
                for run in p.runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(10)
                    if is_header:
                        run.font.bold = True

print("Creating document with bullet points & high side headings...")
new_doc = docx.Document()

# Page setup
sec_cover = new_doc.sections[0]
sec_cover.top_margin = Inches(1.0)
sec_cover.bottom_margin = Inches(1.0)
sec_cover.left_margin = Inches(1.0)
sec_cover.right_margin = Inches(1.0)

# Configure default Normal style (Times New Roman 11pt, 1.15 line spacing)
style_normal = new_doc.styles['Normal']
style_normal.font.name = 'Times New Roman'
style_normal.font.size = Pt(11)
style_normal.font.color.rgb = RGBColor(0, 0, 0)
style_normal.paragraph_format.line_spacing = 1.15
style_normal.paragraph_format.space_after = Pt(4)

# ==========================================
# PAGE 1: COVER PAGE
# ==========================================
p_bar = new_doc.add_paragraph()
p_bar.paragraph_format.space_before = Pt(0)
p_bar.paragraph_format.space_after = Pt(48)
pPr = p_bar._p.get_or_add_pPr()
pBdr = parse_xml(r'<w:pBdr %s><w:bottom w:val="single" w:sz="36" w:space="1" w:color="000000"/></w:pBdr>' % nsdecls('w'))
pPr.append(pBdr)

p_srs = new_doc.add_paragraph()
p_srs.alignment = WD_ALIGN_PARAGRAPH.RIGHT
p_srs.paragraph_format.space_after = Pt(24)
r = p_srs.add_run("Software Requirements\nSpecification")
r.font.name = "Arial"
r.font.size = Pt(26)
r.font.bold = True

p_for = new_doc.add_paragraph()
p_for.alignment = WD_ALIGN_PARAGRAPH.RIGHT
p_for.paragraph_format.space_after = Pt(24)
r = p_for.add_run("for")
r.font.name = "Arial"
r.font.size = Pt(16)
r.font.bold = True

p_proj = new_doc.add_paragraph()
p_proj.alignment = WD_ALIGN_PARAGRAPH.RIGHT
p_proj.paragraph_format.space_after = Pt(48)
r = p_proj.add_run("CourtSetu AI")
r.font.name = "Arial"
r.font.size = Pt(24)
r.font.bold = True

p_ver = new_doc.add_paragraph()
p_ver.alignment = WD_ALIGN_PARAGRAPH.RIGHT
p_ver.paragraph_format.space_after = Pt(36)
r = p_ver.add_run("Version 1.0 approved")
r.font.name = "Arial"
r.font.size = Pt(13)
r.font.bold = True

p_prep = new_doc.add_paragraph()
p_prep.alignment = WD_ALIGN_PARAGRAPH.RIGHT
p_prep.paragraph_format.space_after = Pt(36)
r = p_prep.add_run("Prepared by G Sadwik & Russel Shereef")
r.font.name = "Arial"
r.font.size = Pt(13)
r.font.bold = True

p_org = new_doc.add_paragraph()
p_org.alignment = WD_ALIGN_PARAGRAPH.RIGHT
p_org.paragraph_format.space_after = Pt(36)
r = p_org.add_run("VIT CHENNAI")
r.font.name = "Arial"
r.font.size = Pt(13)
r.font.bold = True

p_date = new_doc.add_paragraph()
p_date.alignment = WD_ALIGN_PARAGRAPH.RIGHT
p_date.paragraph_format.space_after = Pt(100)
r = p_date.add_run("28-07-2026")
r.font.name = "Arial"
r.font.size = Pt(13)
r.font.bold = True

p_copy = new_doc.add_paragraph()
p_copy.alignment = WD_ALIGN_PARAGRAPH.LEFT
r_copy = p_copy.add_run("Copyright © 1999 by Karl E. Wiegers. Permission is granted to use, modify, and distribute this document.")
r_copy.font.name = "Times New Roman"
r_copy.font.size = Pt(9)
r_copy.font.italic = True

# ==========================================
# PAGE 2: TABLE OF CONTENTS & REVISION HISTORY
# ==========================================
sec_toc = new_doc.add_section(docx.enum.section.WD_SECTION.NEW_PAGE)
sec_toc.header.is_linked_to_previous = False

hdr_toc = sec_toc.header
p_hdr = hdr_toc.paragraphs[0]
p_hdr.alignment = WD_ALIGN_PARAGRAPH.RIGHT
p_hdr.paragraph_format.tab_stops.add_tab_stop(Inches(6.5), WD_TAB_ALIGNMENT.RIGHT)
r_h1 = p_hdr.add_run("Software Requirements Specification for CourtSetu AI\tPage ii")
r_h1.font.name = "Times New Roman"
r_h1.font.size = Pt(10)
r_h1.font.italic = True

p_toc_head = new_doc.add_paragraph()
p_toc_head.paragraph_format.space_before = Pt(12)
p_toc_head.paragraph_format.space_after = Pt(12)
r_th = p_toc_head.add_run("Table of Contents")
r_th.font.name = "Times New Roman"
r_th.font.size = Pt(18)
r_th.font.bold = True

p_toc_field = new_doc.add_paragraph()
p_toc_field.paragraph_format.space_after = Pt(6)
fldSimple = parse_xml(r'<w:fldSimple %s w:instr="TOC \o &quot;1-3&quot; \h \z \u"/>' % nsdecls('w'))
p_toc_field._p.append(fldSimple)

toc_entries = [
    ("Table of Contents", 1, "ii"),
    ("Revision History", 1, "ii"),
    ("1. Introduction", 1, "1"),
    ("1.1 Purpose", 2, "1"),
    ("1.2 Document Conventions", 2, "1"),
    ("1.3 Intended Audience and Reading Suggestions", 2, "1"),
    ("1.4 Product Scope", 2, "1"),
    ("1.5 References", 2, "1"),
    ("2. Overall Description", 1, "2"),
    ("2.1 Product Perspective", 2, "2"),
    ("2.2 Product Functions", 2, "2"),
    ("2.3 User Classes and Characteristics", 2, "2"),
    ("2.4 Operating Environment", 2, "2"),
    ("2.5 Design and Implementation Constraints", 2, "2"),
    ("2.6 User Documentation", 2, "2"),
    ("2.7 Assumptions and Dependencies", 2, "3"),
    ("3. External Interface Requirements", 1, "3"),
    ("3.1 User Interfaces", 2, "3"),
    ("3.2 Hardware Interfaces", 2, "3"),
    ("3.3 Software Interfaces", 2, "3"),
    ("3.4 Communications Interfaces", 2, "3"),
    ("4. System Features", 1, "4"),
    ("4.1 User Management", 2, "4"),
    ("4.2 Case Management", 2, "4"),
    ("4.3 AI Services", 2, "4"),
    ("4.4 Lawyer Marketplace", 2, "4"),
    ("4.5 Administration and Analytics", 2, "4"),
    ("5. Other Nonfunctional Requirements", 1, "4"),
    ("5.1 Performance Requirements", 2, "4"),
    ("5.2 Safety Requirements", 2, "5"),
    ("5.3 Security Requirements", 2, "5"),
    ("5.4 Software Quality Attributes", 2, "5"),
    ("5.5 Business Rules", 2, "5"),
    ("6. Other Requirements", 1, "5"),
    ("Appendix A: Glossary", 1, "5"),
    ("Appendix B: Analysis Models", 1, "5"),
    ("Appendix C: References & To Be Determined List", 1, "6"),
    ("Conclusion", 1, "6"),
]

for title, level, page in toc_entries:
    p = new_doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.tab_stops.add_tab_stop(Inches(6.5), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)

    if level == 1:
        p.paragraph_format.left_indent = Inches(0.0)
        p.paragraph_format.space_before = Pt(3)
    elif level == 2:
        p.paragraph_format.left_indent = Inches(0.25)
    elif level == 3:
        p.paragraph_format.left_indent = Inches(0.50)

    r_title = p.add_run(f"{title}\t")
    r_title.font.name = 'Times New Roman'
    r_title.font.size = Pt(10.5)
    if level == 1:
        r_title.font.bold = True
    
    r_page = p.add_run(page)
    r_page.font.name = 'Times New Roman'
    r_page.font.size = Pt(10.5)
    if level == 1:
        r_page.font.bold = True

p_rev_head = new_doc.add_paragraph()
p_rev_head.paragraph_format.space_before = Pt(16)
p_rev_head.paragraph_format.space_after = Pt(8)
r_rh = p_rev_head.add_run("Revision History")
r_rh.font.name = "Times New Roman"
r_rh.font.size = Pt(16)
r_rh.font.bold = True

rev_table = new_doc.add_table(rows=4, cols=4)
rev_headers = ["Name", "Date", "Reason For Changes", "Version"]
for j, text in enumerate(rev_headers):
    rev_table.cell(0, j).text = text

rev_data = [
    ["G Sadwik & Russel Shereef", "28-07-2026", "Initial Draft of Software Requirements Specification for CourtSetu AI", "1.0 approved"],
    ["G Sadwik", "28-07-2026", "Added AI and ML System Requirements & Database Schemas", "1.0.1"],
    ["Russel Shereef", "28-07-2026", "Refined Auction-Based Hiring and Lawyer Verification Requirements", "1.0.2"],
]

for row_idx, data in enumerate(rev_data, start=1):
    for col_idx, cell_value in enumerate(data):
        rev_table.cell(row_idx, col_idx).text = cell_value

style_wiegers_table(rev_table)

# ==========================================
# PAGE 3+: BODY SECTIONS WITH BULLETS & HIGH SIDE HEADINGS
# ==========================================
sec_body = new_doc.add_section(docx.enum.section.WD_SECTION.NEW_PAGE)
sec_body.header.is_linked_to_previous = False

hdr_body = sec_body.header
p_hb = hdr_body.paragraphs[0]
p_hb.alignment = WD_ALIGN_PARAGRAPH.LEFT
p_hb.paragraph_format.tab_stops.add_tab_stop(Inches(6.5), WD_TAB_ALIGNMENT.RIGHT)

r_left = p_hb.add_run("Software Requirements Specification for CourtSetu AI\tPage ")
r_left.font.name = "Times New Roman"
r_left.font.size = Pt(10)
r_left.font.italic = True

fld_page = parse_xml(r'<w:fldSimple %s w:instr="PAGE"/>' % nsdecls('w'))
p_hb._p.append(fld_page)

print("Building body text with bullet points and bold side headings...")

# Known Side Headings / Sub-labels that should be prominent side headings
SIDE_HEADINGS = {
    "User Interface Characteristics", "Client Devices", "Server Infrastructure", "Optional Hardware",
    "Data Exchange", "Communication Protocols", "Internal Communication", "External Communication",
    "User Communication", "Communication Security", "Priority: High", "Priority: Medium",
    "Primary Database Entities", "Database Requirements", "Artificial Intelligence Module",
    "Machine Learning Module", "Privacy and Legal Compliance", "Future Enhancements",
    "Analysis Models Included", "References List"
}

# Key attribute prefixes that should be formatted as bold lead-in bullets (e.g. Availability: ...)
ATTR_PREFIXES = ["Availability:", "Reliability:", "Scalability:", "Maintainability:", "Usability:", "Portability:"]

in_body = False
for p_src in src_doc.paragraphs:
    txt = p_src.text.strip()
    if not txt:
        continue

    if txt in ("Introduction", "1. Introduction", "1.1 Purpose"):
        in_body = True

    if not in_body:
        continue

    if "..........." in txt or txt.startswith("Table of Contents"):
        continue

    # Level 1 Heading (1. Introduction, 2. Overall Description, etc.)
    if re.match(r"^([1-6]\.\s+[A-Z]|Appendix\s+[A-C]|Conclusion|Introduction)", txt):
        if txt == "Introduction":
            txt = "1. Introduction"
        p = new_doc.add_paragraph()
        p.paragraph_format.space_before = Pt(16)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.keep_with_next = True
        r = p.add_run(txt)
        r.font.name = "Times New Roman"
        r.font.size = Pt(16)
        r.font.bold = True
        continue

    # Level 2 Heading (1.1 Purpose, 2.1 Product Perspective, etc.)
    if re.match(r"^[1-6]\.[0-9]+\s+[A-Z]", txt):
        p = new_doc.add_paragraph()
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.keep_with_next = True
        r = p.add_run(txt)
        r.font.name = "Times New Roman"
        r.font.size = Pt(13)
        r.font.bold = True
        continue

    # Level 3 Heading (4.1.1 Description and Priority, etc.)
    if re.match(r"^[1-6]\.[0-9]+\.[0-9]+\s+[A-Z]", txt):
        p = new_doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.keep_with_next = True
        r = p.add_run(txt)
        r.font.name = "Times New Roman"
        r.font.size = Pt(11)
        r.font.bold = True
        continue

    # REQ requirement tags (e.g. REQ-001: ...)
    if txt.startswith("REQ-") or txt.startswith("REQ:"):
        p = new_doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.line_spacing = 1.15

        parts = txt.split(":", 1)
        if len(parts) == 2:
            r_req = p.add_run(parts[0] + ": ")
            r_req.font.name = "Times New Roman"
            r_req.font.size = Pt(11)
            r_req.font.bold = True

            r_text = p.add_run(parts[1].strip())
            r_text.font.name = "Times New Roman"
            r_text.font.size = Pt(11)
        else:
            r = p.add_run(txt)
            r.font.name = "Times New Roman"
            r.font.size = Pt(11)
            r.font.bold = True
        continue

    # High Side Headings (e.g. Client Devices, Server Infrastructure, Priority: High)
    if txt in SIDE_HEADINGS or txt.startswith("Priority:") or (len(txt) < 45 and not txt.endswith(".") and not txt.endswith(";")):
        p = new_doc.add_paragraph()
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.keep_with_next = True
        r = p.add_run(txt)
        r.font.name = "Times New Roman"
        r.font.size = Pt(11.5)
        r.font.bold = True
        r.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
        continue

    # Attribute Lead-in Bullets (e.g. Availability: The platform should...)
    matched_attr = False
    for prefix in ATTR_PREFIXES:
        if txt.startswith(prefix):
            matched_attr = True
            p = new_doc.add_paragraph(style='List Bullet')
            p.paragraph_format.space_after = Pt(3)
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.line_spacing = 1.15

            r_lbl = p.add_run(prefix + " ")
            r_lbl.font.name = "Times New Roman"
            r_lbl.font.size = Pt(11)
            r_lbl.font.bold = True

            r_val = p.add_run(txt[len(prefix):].strip())
            r_val.font.name = "Times New Roman"
            r_val.font.size = Pt(11)
            break

    if matched_attr:
        continue

    # Convert standard list items to BULLET POINTS (•)
    is_list_item = False
    if (
        txt.startswith("•") or txt.startswith("- ") or txt.startswith("▪") or
        txt.endswith(".") and len(txt) < 120 or
        txt.startswith("React.js") or txt.startswith("FastAPI") or txt.startswith("PostgreSQL") or
        txt.startswith("OpenAI") or txt.startswith("Google Maps") or txt.startswith("Razorpay") or
        txt.startswith("OWASP") or txt.startswith("IEEE") or txt.startswith("Karl E.") or
        txt in ("Users", "Lawyers", "Cases", "Case Documents", "Lawyer Proposals (Bids)", "Consultations", "Payments", "Reviews and Ratings", "Notifications", "AI Analysis Reports", "Audit Logs") or
        txt in ("Desktop Computers", "Laptops", "Tablets", "Smartphones (via web browser)", "Cloud-hosted Linux Server", "Multi-core Processor", "Minimum 8 GB RAM", "SSD Storage", "High-speed Internet Connectivity", "HTTPS", "REST API", "JSON", "System Architecture Diagram", "Use Case Diagram", "Entity Relationship (ER) Diagram", "Activity Diagram", "Data Flow Diagram (DFD)")
    ):
        is_list_item = True

    if is_list_item:
        bullet_text = re.sub(r"^[•\-▪]\s*", "", txt)
        p = new_doc.add_paragraph(bullet_text, style='List Bullet')
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.line_spacing = 1.15
        p.paragraph_format.left_indent = Inches(0.25)
        for run in p.runs:
            run.font.name = "Times New Roman"
            run.font.size = Pt(11)
        continue

    # Regular body paragraph
    p = new_doc.add_paragraph(txt)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(4)
    for run in p.runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(11)

print("Copying document tables...")
for t_src in src_doc.tables[1:]:
    t_new = new_doc.add_table(rows=len(t_src.rows), cols=len(t_src.columns))
    for r_idx, row in enumerate(t_src.rows):
        for c_idx, cell in enumerate(row.cells):
            t_new.cell(r_idx, c_idx).text = cell.text.strip()
    style_wiegers_table(t_new)

# Save output
print(f"Saving formatted document to {DST_FILE}...")
new_doc.save(DST_FILE)
try:
    new_doc.save(DST_FILE2)
except Exception as e:
    print(f"Note: DST_FILE2 skipped ({e})")

print("SUCCESS! Document updated with bullet points, high side headings, and clean line spacing!")
